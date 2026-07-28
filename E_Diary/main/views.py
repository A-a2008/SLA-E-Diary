import json
import re
import datetime
import calendar
import logging
from io import BytesIO

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User, Group
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib import messages
from django.core.paginator import Paginator
from django.utils.crypto import get_random_string
from django.db.models import Q
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

from .models import Party1Type, Party2Type, Jurisdiction, CourtLevel, MediationStatus, MediationEntryType, Case, DiaryEntry, CauseListEntry, UserProfile, UserRole, CourtHallNote, Reminder, CourtHallIncharge, SiteSetting
from .constants import COURT_LABELS, BUILDING_LABELS, BUILDING_ORDER, COURT_TO_BUILDING, COURT_HALLS, COURT_HALL_FLOORS
from .services import search_cases, get_latest_entry_data, create_diary_entry, create_case, dispose_case, reinstate_case
from .telegram_utils import send_message

logger = logging.getLogger(__name__)


FAMILY_HALL_ORDER = ['Prl. FC', '1st FC', '2nd FC', '3rd FC', '4th FC', '5th FC', '6th FC']


def _tertiary_sort_key(e):
    court = getattr(e, 'mediation_court', None) or e.case.court
    list_sum = e.mediation_time.hour * 60 + e.mediation_time.minute if e.mediation_time else (e.list_i or 0) + (e.list_ii or 0)
    if court == 'family':
        hall = getattr(e, 'mediation_court_hall', None) or e.case.court_hall
        hi = FAMILY_HALL_ORDER.index(hall) if hall in FAMILY_HALL_ORDER else 999
        return (hi, list_sum)
    return (0, list_sum)


def _get_family_further_dates(date_obj):
    if not date_obj:
        return {}
    from main.models import DiaryEntry
    qs = DiaryEntry.objects.filter(
        case__court='family', next_date__gt=date_obj
    ).values('case__court_hall', 'next_date').distinct().order_by('case__court_hall', 'next_date')
    result = {}
    for row in qs:
        hall = row['case__court_hall']
        if hall not in result:
            result[hall] = []
        result[hall].append(f"{row['next_date'].day}/{row['next_date'].month}")
    return result


def _get_wed_sat_dates(ref_date=None):
    if ref_date is None:
        ref_date = datetime.date.today()
    target_month = ref_date.month + 3
    target_year = ref_date.year
    if target_month > 12:
        target_month -= 12
        target_year += 1
    last_day = calendar.monthrange(target_year, target_month)[1]
    end_date = datetime.date(target_year, target_month, min(ref_date.day, last_day))
    start = ref_date + datetime.timedelta(days=1)

    days_until_wed = (2 - start.weekday()) % 7
    wednesdays = []
    current = start + datetime.timedelta(days=days_until_wed)
    while current <= end_date:
        wednesdays.append(f"{current.day}/{current.month}")
        current += datetime.timedelta(days=7)

    days_until_sat = (5 - start.weekday()) % 7
    saturdays = []
    current = start + datetime.timedelta(days=days_until_sat)
    while current <= end_date:
        saturdays.append(f"{current.day}/{current.month}")
        current += datetime.timedelta(days=7)

    return wednesdays, saturdays


# ── ADMIN / SUPERUSER USER MANAGEMENT ──

@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile') and u.userprofile.role == UserRole.ADMIN or u.is_superuser)
def admin_create_user(request):
    admin_roles = [r for r in UserRole.choices if r[0] != UserRole.ADMIN]
    if request.method == 'POST':
        username = (request.POST.get('username') or '').strip()
        password = request.POST.get('password')
        first_name = (request.POST.get('first_name', '') or '').strip()
        last_name = (request.POST.get('last_name', '') or '').strip()
        role = request.POST.get('role', UserRole.INTERN)
        phone = (request.POST.get('phone', '') or '').strip()
        if User.objects.filter(username=username).exists():
            return render(request, 'registration/admin_create_user.html', {
                'error': 'Username already exists.',
                'role_choices': admin_roles,
            })
        user = User.objects.create_user(username=username, password=password, first_name=first_name, last_name=last_name)
        UserProfile.objects.create(user=user, role=role, phone=phone)
        messages.success(request, f'User "{username}" ({role}) created successfully.')
        return redirect('admin_create_user')
    return render(request, 'registration/admin_create_user.html', {
        'role_choices': admin_roles,
    })


@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile') and u.userprofile.role == UserRole.ADMIN or u.is_superuser)
def manage_users(request):
    if request.user.is_superuser:
        users = User.objects.all().select_related('userprofile').order_by('username')
    else:
        users = User.objects.exclude(is_superuser=True).exclude(userprofile__role=UserRole.ADMIN).select_related('userprofile').order_by('username')
    return render(request, 'registration/manage_users.html', {'users': users})


@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile') and u.userprofile.role == UserRole.ADMIN or u.is_superuser)
def toggle_user_active(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if not request.user.is_superuser and (user.is_superuser or (hasattr(user, 'userprofile') and user.userprofile.role == UserRole.ADMIN)):
        messages.error(request, 'You cannot suspend another admin.')
        return redirect('manage_users')
    user.is_active = not user.is_active
    user.save()
    if hasattr(user, 'userprofile'):
        profile = user.userprofile
        if not user.is_active:
            profile.left_on = datetime.date.today()
        else:
            profile.left_on = None
        profile.save()
    messages.success(request, f'User "{user.username}" {"activated" if user.is_active else "suspended"}.')
    return redirect('manage_users')


@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile') and u.userprofile.role == UserRole.ADMIN or u.is_superuser)
def admin_reset_password(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if not request.user.is_superuser and (user.is_superuser or (hasattr(user, 'userprofile') and user.userprofile.role == UserRole.ADMIN)):
        messages.error(request, 'You cannot reset another admin\'s password.')
        return redirect('manage_users')
    new_password = get_random_string(length=12)
    user.set_password(new_password)
    user.save()
    messages.success(request, f'Password for "{user.username}" reset to: {new_password}')
    return redirect('manage_users')


@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile') and u.userprofile.role == UserRole.ADMIN or u.is_superuser)
def user_detail(request, user_id):
    user = get_object_or_404(User, id=user_id)
    if not request.user.is_superuser and (user.is_superuser or (hasattr(user, 'userprofile') and user.userprofile.role == UserRole.ADMIN)):
        messages.error(request, 'You cannot view another admin\'s details.')
        return redirect('manage_users')
    try:
        profile = user.userprofile
    except UserProfile.DoesNotExist:
        profile = None
    return render(request, 'registration/user_detail.html', {
        'profile_user': user,
        'profile': profile,
    })


@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile') and u.userprofile.role == UserRole.ADMIN or u.is_superuser)
def regenerate_telegram_code(request, user_id):
    user = get_object_or_404(User, id=user_id)
    try:
        profile = user.userprofile
    except UserProfile.DoesNotExist:
        messages.error(request, 'User has no profile.')
        return redirect('manage_users')

    unlink = request.GET.get('unlink')
    if profile.telegram_chat_id and not unlink:
        messages.warning(request, f'"{user.username}" already linked to Telegram. Use "Unlink & Regenerate" first.')
    else:
        if profile.telegram_chat_id:
            profile.telegram_chat_id = None
        profile.telegram_code = UserProfile._generate_code()
        profile.save()
        messages.success(request, f'New Telegram code for "{user.username}": {profile.telegram_code}')

    return redirect('user_detail', user_id=user.id)


@login_required
@user_passes_test(lambda u: hasattr(u, 'userprofile') and u.userprofile.role == UserRole.ADMIN or u.is_superuser)
def toggle_ecourts_access(request, user_id):
    user = get_object_or_404(User, id=user_id)
    try:
        profile = user.userprofile
    except UserProfile.DoesNotExist:
        messages.error(request, 'User has no profile.')
        return redirect('manage_users')
    profile.can_access_ecourts = not profile.can_access_ecourts
    profile.save()
    status = 'granted' if profile.can_access_ecourts else 'revoked'
    messages.success(request, f'eCourts access {status} for "{user.username}".')
    return redirect('user_detail', user_id=user.id)


# ── SUPERUSER PORTAL ──

@login_required
@user_passes_test(lambda u: u.is_superuser)
def super_dashboard(request):
    return render(request, 'registration/super_dashboard.html')


@login_required
@user_passes_test(lambda u: u.is_superuser)
def super_create_admin(request):
    if request.method == 'POST':
        username = (request.POST.get('username') or '').strip()
        password = request.POST.get('password')
        first_name = (request.POST.get('first_name', '') or '').strip()
        last_name = (request.POST.get('last_name', '') or '').strip()
        phone = (request.POST.get('phone', '') or '').strip()
        if User.objects.filter(username=username).exists():
            return render(request, 'registration/super_create_admin.html', {
                'error': 'Username already exists.',
            })
        user = User.objects.create_user(username=username, password=password, first_name=first_name, last_name=last_name)
        user.is_staff = True
        user.save()
        UserProfile.objects.create(user=user, role=UserRole.ADMIN, phone=phone)
        messages.success(request, f'Admin "{username}" created successfully.')
        return redirect('super_dashboard')
    return render(request, 'registration/super_create_admin.html')


def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            next_url = request.GET.get('next', 'home')
            return redirect(next_url)
        else:
            return render(request, 'registration/login.html', {'error': 'Invalid username or password.'})
    return render(request, 'registration/login.html')


@login_required
def logout_view(request):
    logout(request)
    return redirect('home')


@login_required
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)
            messages.success(request, 'Password changed successfully.')
            return redirect('home')
    else:
        form = PasswordChangeForm(request.user)
    return render(request, 'registration/change_password.html', {'form': form})


# ── CASE ──

@login_required
def new_case(request):
    if request.method == 'POST':
        court_level = (request.POST.get('court_level') or '').strip()
        jurisdiction = (request.POST.get('jurisdiction') or '').strip()
        court = (request.POST.get('court') or '').strip()
        court_hall = (request.POST.get('court_hall') or '').strip()
        case_type = (request.POST.get('case_type') or '').strip()
        case_number = (request.POST.get('case_number') or '').strip()
        party_1 = (request.POST.get('party_1') or '').strip()
        party_1_type = (request.POST.get('party_1_type') or '').strip()
        party_2 = (request.POST.get('party_2') or '').strip()
        party_2_type = (request.POST.get('party_2_type') or '').strip()
        representing = (request.POST.get('representing') or '').strip()
        cnr = (request.POST.get('cnr') or '').strip()

        floor = int(request.POST.get('floor') or 0)
        case_year = int(request.POST.get('case_year') or 2024)

        party_1_total = int(request.POST.get('party_1_total') or 1)
        party_2_total = int(request.POST.get('party_2_total') or 1)

        safe_rep = re.sub(r'\W+', '_', representing.lower()).strip('_')
        representing_party_field = f'representing_{safe_rep}_indices'
        representing_parties_list = request.POST.getlist(representing_party_field)
        if representing_parties_list:
            representing_parties = ','.join(representing_parties_list)
        else:
            representing_parties = request.POST.get('representing_parties', '1')

        existing = Case.objects.filter(court=court, case_type=case_type, case_number=case_number, case_year=case_year).first()
        if existing:
            messages.error(request, f'A case with this number already exists: {existing.case_type}/{existing.case_number}/{existing.case_year} — {existing.party_1} vs {existing.party_2}')
            return redirect('diary_entry_case', case_id=existing.id)

        case = create_case(
            jurisdiction=jurisdiction, court_level=court_level, court=court,
            court_hall=court_hall, floor=floor, case_type=case_type,
            case_number=case_number, case_year=case_year, party_1=party_1,
            party_1_type=party_1_type, party_2=party_2, party_2_type=party_2_type,
            representing=representing,
            representing_parties=representing_parties,
            party_1_total=party_1_total,
            party_2_total=party_2_total,
            cnr=cnr,
        )

        CourtHallNote.objects.get_or_create(court=court, court_hall=court_hall, defaults={'note': ''})

        return redirect("diary_entry")
    else:
        import json
        data = {
            'party1_choices': Party1Type.choices,
            'party2_choices': Party2Type.choices,
            'jurisdiction_choices': Jurisdiction.choices,
            'court_level_choices': CourtLevel.choices,
            'court_hall_floors_json': json.dumps(COURT_HALL_FLOORS),
        }
        return render(request, 'main/new_case.html', data)


@login_required
def diary_entry(request):
    query = request.GET.get('q', '').strip()
    court_level = request.GET.get('court_level', '')
    disposed_filter = request.GET.get('disposed', 'active')
    cases = search_cases(query=query, court_level=court_level, disposed=disposed_filter)

    today = datetime.date.today()
    case_list = list(cases)

    def _closest_to_today(*dates):
        valid = [d for d in dates if d is not None]
        if not valid:
            return None
        future = [d for d in valid if d >= today]
        if future:
            return min(future)
        return max(valid)

    def sort_key(case):
        mediation_date = case.mediation_next_date if case.mediation_status in ('referred', 'ongoing') else None
        last_entry = case.diary_entries.order_by('-next_date').first()
        court_date = last_entry.next_date if last_entry else None
        nd = _closest_to_today(mediation_date, court_date)
        if nd is None:
            return (0, datetime.date.min, -case.id)
        if nd >= today:
            return (1, nd, -case.id)
        return (2, nd, -case.id)

    case_list.sort(key=sort_key)

    for case in case_list:
        case.court_display_name = COURT_LABELS.get(case.court, case.court)
        mediation_date = case.mediation_next_date if case.mediation_status in ('referred', 'ongoing') else None
        last_entry = case.diary_entries.order_by('-next_date').first()
        court_date = last_entry.next_date if last_entry else None
        case.next_date = _closest_to_today(mediation_date, court_date)
        case.is_mediation = (case.next_date == mediation_date and mediation_date is not None)
        if case.is_mediation:
            case.prev_date = None
        elif last_entry:
            case.prev_date = last_entry.previous_date
        else:
            case.prev_date = None

    paginator = Paginator(case_list, 20)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    return render(request, 'main/diary_entry.html', {
        'today': today, 'page_obj': page_obj, 'cases': page_obj.object_list,
        'query': query, 'court_labels': COURT_LABELS,
        'court_level_choices': CourtLevel.choices,
        'selected_court_level': court_level, 'selected_disposed': disposed_filter,
    })


@login_required
def diary_entry_case(request, case_id):
    case = get_object_or_404(Case, id=case_id)
    entries = case.diary_entries.all()
    latest_data = get_latest_entry_data(case)

    if request.method == 'POST':
        if request.POST.get('dispose_case') or request.POST.get('reinstate_case'):
            if not request.user.is_superuser:
                return redirect('diary_entry_case', case_id=case.id)
            if request.POST.get('dispose_case'):
                dispose_case(case)
            else:
                reinstate_case(case)
            return redirect('diary_entry_case', case_id=case.id)

    court_hall_notes = CourtHallNote.objects.filter(
        court=case.court, court_hall=case.court_hall
    )

    ctx = {
        'case': case, 'entries': entries, 'court_labels': COURT_LABELS,
        'latest_data': latest_data, 'court_display': COURT_LABELS.get(case.court, case.court),
        'court_hall_notes': court_hall_notes,
        'mediation_statuses': MediationStatus,
    }

    if request.user.is_superuser or request.user.groups.filter(name='payments').exists():
        try:
            from payments.models import (
                ChargeType, CaseChargeAmount, EntryClassification,
                CasePricing
            )
            from payments.services import get_applicable_charge_types
            pricing = CasePricing.objects.filter(case=case).first()
            charge_types = get_applicable_charge_types(case) if pricing else []
            charge_amounts = dict(
                CaseChargeAmount.objects.filter(case_pricing=pricing)
                .values_list('charge_type_id', 'amount')
            ) if pricing else {}
            classifications = {}
            for entry in entries:
                try:
                    cls = entry.classification
                    selected = set(cls.charge_items.filter(
                        charge_type__isnull=False
                    ).values_list('charge_type_id', flat=True))
                    classifications[entry.id] = selected
                except EntryClassification.DoesNotExist:
                    classifications[entry.id] = set()
            ctx['show_payments_ui'] = True
            ctx['payments_charge_types'] = charge_types
            ctx['payments_charge_amounts'] = charge_amounts
            ctx['payments_entry_classifications'] = classifications
        except Exception:
            pass

    return render(request, 'main/diary_entry_case.html', ctx)


@login_required
def refer_to_mediation(request, case_id):
    case = get_object_or_404(Case, id=case_id)

    if request.method == 'POST':
        mediation_date = request.POST.get('mediation_date')
        notes = request.POST.get('notes', '').strip()

        if mediation_date:
            try:
                mediation_date_obj = datetime.datetime.strptime(mediation_date, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Invalid date.')
                return redirect('diary_entry_case', case_id=case.id)

            case.mediation_status = MediationStatus.REFERRED
            case.mediation_next_date = mediation_date_obj
            case.save()

            business_text = 'Case referred to Karnataka Mediation Centre.'
            if notes:
                business_text += f'\n\nNotes: {notes}'

            create_diary_entry(
                case=case,
                entry_type='mediation',
                previous_date=datetime.date.today(),
                court='Karnataka Mediation Centre',
                court_hall='Mediation',
                floor=0,
                case_number_display=f"{case.case_type}/{case.case_number}/{case.case_year}",
                representing=case.representing,
                representing_parties=case.representing_parties,
                party_1_total=case.party_1_total,
                party_2_total=case.party_2_total,
                stage='Mediation',
                business=business_text,
                next_date=mediation_date_obj,
                advocate=request.user,
            )

            messages.success(request, f'Case referred to mediation. Next mediation date: {mediation_date}')
        else:
            messages.error(request, 'Please provide a mediation date.')

        return redirect('diary_entry_case', case_id=case.id)

    return render(request, 'main/refer_to_mediation.html', {
        'case': case, 'court_display': COURT_LABELS.get(case.court, case.court),
        'mediation_statuses': MediationStatus,
    })


@login_required
def update_mediation_status(request, case_id):
    case = get_object_or_404(Case, id=case_id)

    if request.method == 'POST':
        new_status = request.POST.get('mediation_status')
        mediation_date = request.POST.get('mediation_date')

        if new_status and new_status in [s.value for s in MediationStatus]:
            case.mediation_status = new_status
            if mediation_date:
                try:
                    case.mediation_next_date = datetime.datetime.strptime(mediation_date, '%Y-%m-%d').date()
                except ValueError:
                    pass
            elif new_status in (MediationStatus.SETTLED, MediationStatus.FAILED):
                case.mediation_next_date = None
            case.save()

            if new_status in (MediationStatus.SETTLED, MediationStatus.FAILED):
                label = dict(MediationStatus.choices).get(new_status, new_status)
                create_diary_entry(
                    case=case,
                    entry_type='mediation',
                    previous_date=datetime.date.today(),
                    court='Karnataka Mediation Centre',
                    court_hall='Mediation',
                    floor=0,
                    case_number_display=f"{case.case_type}/{case.case_number}/{case.case_year}",
                    representing=case.representing,
                    representing_parties=case.representing_parties,
                    party_1_total=case.party_1_total,
                    party_2_total=case.party_2_total,
                    stage='Mediation',
                    business=f'Mediation {label.lower()}.',
                    next_date=case.mediation_next_date or datetime.date.today(),
                    advocate=request.user,
                )

            messages.success(request, f'Mediation status updated to "{dict(MediationStatus.choices).get(new_status, new_status)}".')
        else:
            messages.error(request, 'Invalid status.')

        return redirect('diary_entry_case', case_id=case.id)


@login_required
def create_execution_case(request, case_id):
    original_case = get_object_or_404(Case, id=case_id)

    # Determine default swap: if advocate represents party 2, swap so client becomes party 1
    default_swap = original_case.representing_parties == '2'

    if request.method == 'POST':
        case_number = (request.POST.get('case_number') or '').strip()
        case_year = (request.POST.get('case_year') or '').strip()
        party_1 = (request.POST.get('party_1') or original_case.party_1).strip()
        party_2 = (request.POST.get('party_2') or original_case.party_2).strip()
        representing = (request.POST.get('representing') or original_case.representing).strip()

        try:
            case = Case.objects.create(
                jurisdiction=original_case.jurisdiction,
                court_level='district',
                court=original_case.court,
                court_hall=original_case.court_hall,
                floor=original_case.floor,
                case_type='EX',
                case_number=case_number,
                case_year=int(case_year) if case_year else datetime.date.today().year,
                party_1=party_1,
                party_1_type='Decree Holder',
                party_2=party_2,
                party_2_type='Judgement Debtor',
                representing=representing,
                representing_parties='1',
                related_case=original_case,
            )
            messages.success(request, f'Execution case created: EX/{case_number}/{case_year}')
            return redirect('diary_entry_case', case_id=case.id)
        except Exception as e:
            messages.error(request, f'Error creating case: {e}')

    today = datetime.date.today()

    # Build initial values, applying swap if advocate represented party 2
    party_1 = original_case.party_2 if default_swap else original_case.party_1
    party_2 = original_case.party_1 if default_swap else original_case.party_2

    court_label = COURT_LABELS.get(original_case.court, original_case.court)

    return render(request, 'main/create_execution_case.html', {
        'original_case': original_case,
        'party_1': party_1,
        'party_2': party_2,
        'default_swap': default_swap,
        'court_label': court_label,
        'today': today,
    })


@login_required
def add_business(request, case_id):
    case = get_object_or_404(Case, id=case_id)

    latest_data = get_latest_entry_data(case)

    if request.method == 'POST':
        previous_date = (request.POST.get('previous_date') or '').strip()

        # Validate that previous_date matches the last business entry's next_date
        last_entry = DiaryEntry.objects.filter(case=case, entry_type='business').order_by('-next_date').first()
        if last_entry and previous_date:
            try:
                prev_date_obj = datetime.datetime.strptime(previous_date, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, 'Invalid date.')
                return redirect('add_business', case_id=case.id)
            if prev_date_obj != last_entry.next_date:
                messages.error(request, f'The court date must be {last_entry.next_date}. You can only add a new entry on the scheduled court date.')
                return redirect('diary_entry_case', case_id=case.id)

        # If case is in mediation, mark it ongoing when a court entry is added
        if case.mediation_status == 'referred':
            case.mediation_status = 'ongoing'
            case.mediation_next_date = None
            case.save()
        elif case.mediation_status == 'ongoing':
            case.save()
        court = COURT_LABELS.get(case.court, case.court)
        court_hall = case.court_hall
        floor = case.floor
        case_number_display = f"{case.case_type}/{case.case_number}/{case.case_year}"
        representing = (request.POST.get('representing') or case.representing).strip()
        stage = (request.POST.get('stage') or '').strip()
        business = (request.POST.get('business') or '').strip()
        next_date = (request.POST.get('next_date') or '').strip()

        safe_rep = re.sub(r'\W+', '_', representing.lower()).strip('_')
        representing_party_field = f'representing_{safe_rep}_indices'
        representing_parties_list = request.POST.getlist(representing_party_field)
        if representing_parties_list:
            representing_parties = ','.join(representing_parties_list)
        else:
            representing_parties = request.POST.get('representing_parties', case.representing_parties)

        party_1_total = int(request.POST.get('party_1_total') or case.party_1_total)
        party_2_total = int(request.POST.get('party_2_total') or case.party_2_total)

        entry = create_diary_entry(
            case=case, previous_date=previous_date, court=court,
            court_hall=court_hall, floor=floor,
            case_number_display=case_number_display, representing=representing,
            representing_parties=representing_parties,
            party_1_total=party_1_total,
            party_2_total=party_2_total,
            stage=stage, business=business, next_date=next_date,
            advocate=request.user,
        )

        from .ecourts_integration import summarize_business
        summary = summarize_business(
            business, entry.ecourts_business or '', case=case
        )
        if summary:
            entry.business_summary = summary
            entry.save(update_fields=['business_summary'])

        if request.POST.get('needs_reminder'):
            reminder_task = request.POST.get('reminder_task', '').strip()
            reminder_start_on = request.POST.get('reminder_start_on')
            reminder_frequency = request.POST.get('reminder_frequency', 'daily')
            reminder_ramp_up = request.POST.get('reminder_ramp_up') == '1'
            if reminder_task and reminder_start_on:
                Reminder.objects.create(
                    diary_entry=entry,
                    task=reminder_task,
                    start_on=reminder_start_on,
                    frequency=reminder_frequency,
                    ramp_up=reminder_ramp_up,
                )

        return redirect('diary_entry_case', case_id=case.id)

    return render(request, 'main/add_business.html', {
        'case': case, 'latest_data': latest_data,
        'court_display': COURT_LABELS.get(case.court, case.court),
        'today': datetime.date.today(),
    })


@login_required
def edit_business(request, entry_id):
    entry = get_object_or_404(DiaryEntry, id=entry_id)

    if request.method == 'POST':
        entry.previous_date = (request.POST.get('previous_date') or '').strip()
        if entry.entry_type != 'mediation':
            entry.court = COURT_LABELS.get(entry.case.court, entry.case.court)
            entry.court_hall = entry.case.court_hall
            entry.floor = entry.case.floor
        entry.case_number_display = f"{entry.case.case_type}/{entry.case.case_number}/{entry.case.case_year}"
        entry.representing = (request.POST.get('representing') or entry.case.representing).strip()
        entry.stage = (request.POST.get('stage') or '').strip()
        entry.business = (request.POST.get('business') or '').strip()
        entry.next_date = (request.POST.get('next_date') or '').strip()

        safe_rep = re.sub(r'\W+', '_', entry.representing.lower()).strip('_')
        representing_party_field = f'representing_{safe_rep}_indices'
        representing_parties_list = request.POST.getlist(representing_party_field)
        if representing_parties_list:
            entry.representing_parties = ','.join(representing_parties_list)

        party_1_total = request.POST.get('party_1_total')
        if party_1_total:
            entry.party_1_total = int(party_1_total)
        party_2_total = request.POST.get('party_2_total')
        if party_2_total:
            entry.party_2_total = int(party_2_total)

        entry.save()

        if request.POST.get('needs_reminder'):
            reminder_task = request.POST.get('reminder_task', '').strip()
            reminder_start_on = request.POST.get('reminder_start_on')
            reminder_frequency = request.POST.get('reminder_frequency', 'daily')
            reminder_ramp_up = request.POST.get('reminder_ramp_up') == '1'
            if reminder_task and reminder_start_on:
                reminder, created = Reminder.objects.get_or_create(
                    diary_entry=entry,
                    defaults={
                        'task': reminder_task,
                        'start_on': reminder_start_on,
                        'frequency': reminder_frequency,
                        'ramp_up': reminder_ramp_up,
                    }
                )
                if not created:
                    reminder.task = reminder_task
                    reminder.start_on = reminder_start_on
                    reminder.frequency = reminder_frequency
                    reminder.ramp_up = reminder_ramp_up
                    reminder.completed = False
                    reminder.save()
        else:
            Reminder.objects.filter(diary_entry=entry).delete()

        return redirect('diary_entry_case', case_id=entry.case.id)

    reminder = entry.reminders.first()
    return render(request, 'main/edit_business.html', {
        'entry': entry,
        'case': entry.case,
        'court_display': COURT_LABELS.get(entry.case.court, entry.case.court),
        'reminder': reminder,
        'today': datetime.date.today(),
    })


# ── SEARCH ──

@login_required
def case_search(request):
    query = request.GET.get('q', '').strip()
    court_level = request.GET.get('court_level', '')
    disposed_filter = request.GET.get('disposed', '')
    cases = search_cases(query=query, court_level=court_level, disposed=disposed_filter)

    case_list = list(cases)

    def sort_key(case):
        last_entry = case.diary_entries.order_by('-previous_date').first()
        if last_entry is None:
            return (0, -case.id)
        return (1, last_entry.previous_date.isoformat(), -case.id)

    case_list.sort(key=sort_key, reverse=True)

    today = datetime.date.today()

    def _closest_to_today(*dates):
        valid = [d for d in dates if d is not None]
        if not valid:
            return None
        future = [d for d in valid if d >= today]
        if future:
            return min(future)
        return max(valid)

    for case in case_list:
        case.court_display_name = COURT_LABELS.get(case.court, case.court)
        mediation_date = case.mediation_next_date if case.mediation_status in ('referred', 'ongoing') else None
        last_entry = case.diary_entries.order_by('-next_date').first()
        court_date = last_entry.next_date if last_entry else None
        case.next_date = _closest_to_today(mediation_date, court_date)
        if last_entry:
            case.prev_date = last_entry.previous_date
        else:
            case.prev_date = None
    return render(request, 'main/search_cases.html', {
        'today': datetime.date.today(), 'cases': case_list, 'query': query, 'court_labels': COURT_LABELS,
        'court_level_choices': CourtLevel.choices,
        'selected_court_level': court_level, 'selected_disposed': disposed_filter,
    })


# ── CASE EXPORT ──

@login_required
def case_export_docx(request, case_id):
    case = get_object_or_404(Case, id=case_id)
    entries = case.diary_entries.all()

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(f'{case.case_type} {case.case_number}/{case.case_year}', 0)
    p = doc.add_paragraph()
    run1 = p.add_run(case.party_1)
    run1.bold = case.represents_party_1
    p.add_run(' vs ')
    run2 = p.add_run(case.party_2)
    run2.bold = case.represents_party_2
    doc.add_paragraph(f'Party Types: {case.party_1_type} / {case.party_2_type}   |   Representing: {case.representing}')
    doc.add_paragraph(f'Court: {COURT_LABELS.get(case.court, case.court)}, Hall: {case.court_hall}, Floor: {case.floor}')
    doc.add_paragraph(f'Jurisdiction: {case.get_jurisdiction_display()}   |   Level: {case.get_court_level_display()}')
    status = 'DISPOSED' if case.disposed else 'ACTIVE'
    doc.add_paragraph(f'Status: {status}')
    doc.add_paragraph('')

    if entries:
        doc.add_heading('Business History', level=1)
        for entry in entries:
            doc.add_heading(f'{entry.previous_date.strftime("%d %b %Y")}  →  {entry.next_date.strftime("%d %b %Y")}', level=2)
            doc.add_paragraph(f'Stage: {entry.stage}')
            doc.add_paragraph(entry.business)
            doc.add_paragraph(f'Entered by: {entry.advocate.get_full_name() or entry.advocate.username if entry.advocate else "—"} on {entry.created_at.strftime("%d %b %Y %I:%M %p")}')
            doc.add_paragraph('')
    else:
        doc.add_paragraph('No diary entries yet.')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f'{case.case_type}_{case.case_number}_{case.case_year}.docx'.replace('/', '_')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)
    return response


@login_required
def case_export_pdf(request, case_id):
    case = get_object_or_404(Case, id=case_id)
    entries = case.diary_entries.all()

    from io import BytesIO
    from django.template.loader import render_to_string
    from weasyprint import HTML

    html = render_to_string('main/case_export_pdf.html', {
        'case': case, 'entries': entries, 'court_labels': COURT_LABELS,
        'court_display': COURT_LABELS.get(case.court, case.court),
    })
    pdf_file = BytesIO()
    HTML(string=html).write_pdf(pdf_file)
    pdf_file.seek(0)

    filename = f'{case.case_type}_{case.case_number}_{case.case_year}.pdf'.replace('/', '_')
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ── CAUSE LIST ──

@login_required
def cause_list(request):
    date_str = request.GET.get('date', '')
    entries = DiaryEntry.objects.none()
    errors = []

    if request.method == 'POST':
        date_str = request.POST.get('date', '')
        updated = 0
        entry_ids = set()
        for key in request.POST:
            if key.startswith('list_i_'):
                entry_ids.add(key.replace('list_i_', ''))
            elif key.startswith('list_ii_'):
                entry_ids.add(key.replace('list_ii_', ''))
            elif key.startswith('mediation_time_'):
                entry_ids.add(key.replace('mediation_time_', ''))

        for eid in sorted(entry_ids):
            list_i_val = request.POST.get(f'list_i_{eid}', '').strip()
            list_ii_val = request.POST.get(f'list_ii_{eid}', '').strip()
            mediation_time_val = request.POST.get(f'mediation_time_{eid}', '').strip()

            if list_ii_val and not list_i_val:
                errors.append(f'Case #{eid}: List II cannot be entered without List I.')
                continue

            try:
                case = Case.objects.get(id=eid)
            except Case.DoesNotExist:
                continue

            cl_entry, created = CauseListEntry.objects.get_or_create(
                date=datetime.datetime.strptime(date_str, '%Y-%m-%d').date(),
                case=case,
            )

            if mediation_time_val:
                try:
                    cl_entry.mediation_time = datetime.datetime.strptime(mediation_time_val, '%H:%M').time()
                    cl_entry.list_i = None
                    cl_entry.list_ii = None
                    cl_entry.save()
                    updated += 1
                except ValueError:
                    errors.append(f'Case #{eid}: Invalid time format (use HH:MM).')
            elif list_i_val:
                try:
                    cl_entry.list_i = int(list_i_val)
                    cl_entry.list_ii = int(list_ii_val) if list_ii_val else 0
                    cl_entry.mediation_time = None
                    cl_entry.save()
                    updated += 1
                except ValueError:
                    errors.append(f'Case #{eid}: Invalid number.')
            else:
                if cl_entry.list_i is not None or cl_entry.list_ii is not None or cl_entry.mediation_time is not None:
                    cl_entry.list_i = None
                    cl_entry.list_ii = None
                    cl_entry.mediation_time = None
                    cl_entry.save()
                    updated += 1

        if not errors and updated:
            messages.success(request, f'{updated} cause list number(s) updated.')

        # Save incharge checkboxes
        try:
            date_obj_incharge = datetime.datetime.strptime(date_str, '%Y-%m-%d').date()
        except ValueError:
            date_obj_incharge = None
        if date_obj_incharge:
            incharge_keys = set()
            for key in request.POST:
                if key.startswith('incharge_'):
                    court_hall_key = key[len('incharge_'):]
                    incharge_keys.add(court_hall_key)
            CourtHallIncharge.objects.filter(date=date_obj_incharge).delete()
            for ch_key in incharge_keys:
                parts = ch_key.split('__', 1)
                if len(parts) == 2:
                    CourtHallIncharge.objects.create(
                        date=date_obj_incharge,
                        court=parts[0],
                        court_hall=parts[1],
                        is_incharge=True,
                    )

        return redirect(f'{request.path}?date={date_str}')

    if date_str:
        try:
            date_obj = datetime.datetime.strptime(date_str, '%Y-%m-%d').date()
            cl_entries = CauseListEntry.objects.filter(date=date_obj).select_related('case')
            cases_with_cl = {e.case_id for e in cl_entries}

            all_cases = Case.objects.filter(
                diary_entries__next_date=date_obj
            ).distinct()

            for case in all_cases:
                if case.id not in cases_with_cl:
                    CauseListEntry.objects.create(date=date_obj, case=case)

            entries = CauseListEntry.objects.filter(date=date_obj).select_related('case')

            court_order = [
                'city_civil', 'small_causes', 'city_civil_rural', 'dist_session_rural',
                'prl_senior_rural', 'prl_junior_rural',
                'cmm', 'mmtc', 'mmtc_mayo', 'cjm_cmm_rural', 'vacation_court_rural',
                'family',
                'vacation_bench_family',
                'consumer', 'urban_consumer',
                'commercial', 'commercial_court_rural',
                'mayo_hall',
                'labour_court_rural', 'labour_court_urban', 'senior_anekal', 'junior_anekal', 'hosakote',
                'devanahalli', 'doddaballapur', 'nelamangala', 'kr_puram',
                'high_court_karnataka', 'supreme_court_india',
            ]

            diary_entries_for_stage = DiaryEntry.objects.filter(
                next_date=date_obj
            ).values('case_id', 'stage', 'entry_type', 'court', 'court_hall')
            stage_by_case = {}
            mediation_court_by_case = {}
            for de in diary_entries_for_stage:
                stage_by_case[de['case_id']] = de['stage']
                if de['entry_type'] == 'mediation':
                    mc = de['court']
                    mh = de['court_hall']
                    if mc not in COURT_TO_BUILDING:
                        mc = 'Karnataka Mediation Centre'
                        mh = 'Mediation'
                    mediation_court_by_case[de['case_id']] = (mc, mh)
            for e in entries:
                e.stage = stage_by_case.get(e.case.id, '')
                if e.case.id in mediation_court_by_case:
                    e.mediation_court, e.mediation_court_hall = mediation_court_by_case[e.case.id]

            entries = sorted(entries, key=lambda e: (
                BUILDING_ORDER.index(COURT_TO_BUILDING.get(
                    getattr(e, 'mediation_court', None) or e.case.court, 'other'))
                    if (getattr(e, 'mediation_court', None) or e.case.court) in COURT_TO_BUILDING else 999,
                court_order.index(e.case.court) if e.case.court in court_order else 999,
                _tertiary_sort_key(e),
            ))
            for sl, e in enumerate(entries, 1):
                e.sl_no = sl

        except ValueError:
            pass

    court_halls_on_date = []
    seen_halls = set()
    if date_str:
        for e in entries:
            court = getattr(e, 'mediation_court', None) or e.case.court
            hall = getattr(e, 'mediation_court_hall', None) or e.case.court_hall
            key = (court, hall)
            if key not in seen_halls:
                seen_halls.add(key)
                court_halls_on_date.append(key)
    court_hall_notes = dict()
    for n in CourtHallNote.objects.all():
        key = f"{n.court}__{n.court_hall}"
        court_hall_notes[key] = n.note

    court_hall_incharges = set()
    if date_str and date_obj:
        for chi in CourtHallIncharge.objects.filter(date=date_obj, is_incharge=True):
            court_hall_incharges.add(f"{chi.court}__{chi.court_hall}")

    unique_hall_keys = []
    seen = set()
    for e in entries:
        court = getattr(e, 'mediation_court', None) or e.case.court
        court_hall = getattr(e, 'mediation_court_hall', None) or e.case.court_hall
        key = f"{court}__{court_hall}"
        if key in court_hall_notes and key not in seen:
            seen.add(key)
            unique_hall_keys.append((court, court_hall))

    building_groups = []
    if entries:
        from itertools import groupby
        def effective_building_code(e):
            court = getattr(e, 'mediation_court', None) or e.case.court
            return BUILDING_ORDER.index(COURT_TO_BUILDING.get(court, 'other')) if court in COURT_TO_BUILDING else 999
        for bldg_code, group in groupby(entries, key=effective_building_code):
            actual_code = BUILDING_ORDER[bldg_code] if bldg_code < len(BUILDING_ORDER) else 'other'
            building_groups.append((actual_code, list(group)))

    wednesdays, saturdays = _get_wed_sat_dates(date_obj if date_str else None)
    family_further_dates = _get_family_further_dates(date_obj if date_str else None)

    return render(request, 'main/cause_list.html', {
        'entries': entries, 'building_groups': building_groups,
        'date_str': date_str, 'date_obj': date_obj if date_str else None,
        'court_labels': COURT_LABELS, 'errors': errors,
        'court_hall_notes': court_hall_notes,
        'unique_hall_keys': unique_hall_keys,
        'court_hall_incharges': court_hall_incharges,
        'court_halls_on_date': court_halls_on_date,
        'wednesdays': wednesdays,
        'saturdays': saturdays,
        'family_further_dates': family_further_dates,
    })


@login_required
def cause_list_docx(request):
    date_str = request.GET.get('date', '')
    if not date_str:
        return HttpResponse('No date provided.', status=400)

    try:
        date_obj = datetime.datetime.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        return HttpResponse('Invalid date.', status=400)

    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    entries = CauseListEntry.objects.filter(date=date_obj).select_related('case')

    court_order = [
        'city_civil', 'small_causes', 'city_civil_rural', 'dist_session_rural',
        'prl_senior_rural', 'prl_junior_rural',
        'cmm', 'mmtc', 'mmtc_mayo', 'cjm_cmm_rural', 'vacation_court_rural',
        'family',
        'vacation_bench_family',
        'consumer', 'urban_consumer',
        'commercial', 'commercial_court_rural',
        'mayo_hall',
        'labour_court_rural', 'labour_court_urban', 'senior_anekal', 'junior_anekal', 'hosakote',
        'devanahalli', 'doddaballapur', 'nelamangala', 'kr_puram',
        'high_court_karnataka', 'supreme_court_india',
    ]
    diary_entries_for_stage = DiaryEntry.objects.filter(
        next_date=date_obj
    ).values('case_id', 'stage', 'entry_type', 'court', 'court_hall')
    stage_by_case = {}
    mediation_court_by_case = {}
    for de in diary_entries_for_stage:
        stage_by_case[de['case_id']] = de['stage']
        if de['entry_type'] == 'mediation':
            mc = de['court']
            mh = de['court_hall']
            if mc not in COURT_TO_BUILDING:
                mc = 'Karnataka Mediation Centre'
                mh = 'Mediation'
            mediation_court_by_case[de['case_id']] = (mc, mh)
    for e in entries:
        e.stage = stage_by_case.get(e.case.id, '')
        if e.case.id in mediation_court_by_case:
            e.mediation_court, e.mediation_court_hall = mediation_court_by_case[e.case.id]

    entries = sorted(entries, key=lambda e: (
        BUILDING_ORDER.index(COURT_TO_BUILDING.get(
            getattr(e, 'mediation_court', None) or e.case.court, 'other'))
            if (getattr(e, 'mediation_court', None) or e.case.court) in COURT_TO_BUILDING else 999,
        court_order.index(e.case.court) if e.case.court in court_order else 999,
        _tertiary_sort_key(e),
    ))
    for sl, e in enumerate(entries, 1):
        e.sl_no = sl

    court_hall_incharges = set()
    for chi in CourtHallIncharge.objects.filter(date=date_obj, is_incharge=True):
        court_hall_incharges.add(f"{chi.court}__{chi.court_hall}")

    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Helvetica'
    normal_style.font.size = Pt(10)

    def set_cell_shading(cell, color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_width(cell, width):
        cell.width = width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width.emu / 635)}" w:type="dxa"/>')
            tcPr.append(tcW)
        else:
            tcW.set(qn('w:w'), str(int(width.emu / 635)))
            tcW.set(qn('w:type'), 'dxa')

    def set_run_font(run, size=Pt(10)):
        run.font.name = 'Helvetica'
        run.font.size = size

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f'Cause List — {date_obj.strftime("%d %B %Y")}')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Helvetica'
    title_para.paragraph_format.space_after = Pt(0)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(date_obj.strftime("%A"))
    sub_run.bold = True
    sub_run.font.size = Pt(12)
    sub_run.font.name = 'Helvetica'
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub_para.paragraph_format.space_after = Pt(6)

    wednesdays, saturdays = _get_wed_sat_dates(date_obj)
    family_further_dates = _get_family_further_dates(date_obj)
    if wednesdays or saturdays:
        wed_para = doc.add_paragraph()
        wed_run = wed_para.add_run('Wednesdays: ')
        wed_run.bold = True
        wed_run.font.size = Pt(9)
        wed_run.font.name = 'Helvetica'
        wed_val = wed_para.add_run(', '.join(wednesdays))
        wed_val.font.size = Pt(9)
        wed_val.font.name = 'Helvetica'
        wed_para.paragraph_format.space_after = Pt(2)

        sat_para = doc.add_paragraph()
        sat_run = sat_para.add_run('Saturdays: ')
        sat_run.bold = True
        sat_run.font.size = Pt(9)
        sat_run.font.name = 'Helvetica'
        sat_val = sat_para.add_run(', '.join(saturdays))
        sat_val.font.size = Pt(9)
        sat_val.font.name = 'Helvetica'
        sat_para.paragraph_format.space_after = Pt(8)

    from itertools import groupby
    current_building = None

    def _add_docx_data_row(table, entry, col_widths):
        row = table.add_row().cells
        case_num = f"{entry.case.case_type}/{entry.case.case_number}/{entry.case.case_year}"
        effective_hall = getattr(entry, 'mediation_court_hall', None) or entry.case.court_hall
        effective_court = getattr(entry, 'mediation_court', None) or entry.case.court
        effective_court_label = COURT_LABELS.get(effective_court, effective_court)
        bldg_code = COURT_TO_BUILDING.get(effective_court, '')
        if bldg_code == 'mediation_centre':
            data = [str(entry.sl_no), str(entry.case.floor), None, None, entry.case.representing, entry.stage or '—',
                    entry.mediation_time.strftime('%I:%M %p') if entry.mediation_time else '—']
        else:
            cause_list_nos = f"List I: {entry.list_i or '—'}\nList II: {entry.list_ii or '—'}"
            data = [str(entry.sl_no), str(entry.case.floor), None, None, entry.case.representing, entry.stage or '—', cause_list_nos]
        for i, val in enumerate(data):
            cell = row[i]
            set_cell_width(cell, col_widths[i])
            if val is None:
                p = cell.paragraphs[0]
                p.clear()
                if i == 2:
                    run_court = p.add_run(f"{effective_court_label}\n")
                    set_run_font(run_court)
                    run_hall = p.add_run(effective_hall)
                    run_hall.bold = True
                    set_run_font(run_hall)
                elif i == 3:
                    run0 = p.add_run(f"{case_num}\n")
                    run0.bold = True
                    set_run_font(run0)
                    run1 = p.add_run(entry.case.party_1)
                    run1.bold = entry.case.represents_party_1
                    set_run_font(run1)
                    run_vs = p.add_run(' vs ')
                    set_run_font(run_vs)
                    run2 = p.add_run(entry.case.party_2)
                    run2.bold = entry.case.represents_party_2
                    set_run_font(run2)
            else:
                cell.text = ''
                p = cell.paragraphs[0]
                if i in (0, 1):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                set_run_font(run)

    def _add_docx_new_table(doc, bldg_name):
        bldg_para = doc.add_paragraph()
        bldg_para.paragraph_format.space_before = Pt(8)
        bldg_para.paragraph_format.space_after = Pt(2)
        pPr = bldg_para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="dddddd"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        bldg_run = bldg_para.add_run(bldg_name)
        bldg_run.bold = True
        bldg_run.font.size = Pt(13)
        bldg_run.font.name = 'Helvetica'
        bldg_run.font.color.rgb = RGBColor(0xc4, 0x45, 0x69)

        table = doc.add_table(rows=1, cols=7)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)

        existing_tblW = tblPr.find(qn('w:tblW'))
        if existing_tblW is not None:
            tblPr.remove(existing_tblW)
        tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(Cm(18.2).emu / 635)}" w:type="dxa"/>'))

        existing_layout = tblPr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tblPr.remove(existing_layout)
        tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="cccccc"/>'
            f'</w:tblBorders>'
        ))
        return table

    def _add_docx_further_dates_row(table, dates, col_widths):
        row = table.add_row().cells
        merged = row[0].merge(row[-1])
        p = merged.paragraphs[0]
        p.clear()
        run_label = p.add_run('Further dates: ')
        run_label.bold = True
        set_run_font(run_label)
        dates_text = ', '.join(dates) if dates else '—'
        run_val = p.add_run(dates_text)
        set_run_font(run_val)

    i = 0
    while i < len(entries):
        entry = entries[i]
        effective_court = getattr(entry, 'mediation_court', None) or entry.case.court
        bldg_code = COURT_TO_BUILDING.get(effective_court, '')
        bldg_name = BUILDING_LABELS.get(bldg_code, COURT_LABELS.get(effective_court, effective_court))

        if bldg_code == 'family_court':
            if bldg_name != current_building:
                current_building = bldg_name
                table = _add_docx_new_table(doc, bldg_name)
                is_mediation = False
                hdr = table.rows[0].cells
                headers = ['Sl No.', 'Floor', 'Court Hall', 'Case & Parties', 'Representing', 'Stage', 'Cause List']
                col_widths = [Cm(1.5), Cm(1.5), Cm(3.0), Cm(4.9), Cm(2.8), Cm(2.2), Cm(2.3)]
                for j, h in enumerate(headers):
                    set_cell_width(hdr[j], col_widths[j])
                    hdr[j].text = ''
                    p = hdr[j].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(h)
                    run.bold = True
                    set_run_font(run)
                    set_cell_shading(hdr[j], 'f5f5f5')

            # Collect all consecutive family entries
            family_entries = []
            while i < len(entries):
                e = entries[i]
                ec = getattr(e, 'mediation_court', None) or e.case.court
                bc = COURT_TO_BUILDING.get(ec, '')
                if bc != 'family_court':
                    break
                family_entries.append(e)
                i += 1

            # Group by hall and render
            for hall, hall_group in groupby(family_entries, key=lambda e: getattr(e, 'mediation_court_hall', None) or e.case.court_hall):
                hall_list = list(hall_group)
                for entry_h in hall_list:
                    _add_docx_data_row(table, entry_h, col_widths)
                fwd_dates = family_further_dates.get(hall, [])
                _add_docx_further_dates_row(table, fwd_dates, col_widths)
            continue

        else:
            if bldg_name != current_building:
                current_building = bldg_name
                table = _add_docx_new_table(doc, bldg_name)
                is_mediation = (bldg_code == 'mediation_centre')
                hdr = table.rows[0].cells
                headers = ['Sl No.', 'Floor', 'Court Hall', 'Case & Parties', 'Representing', 'Stage', 'Time' if is_mediation else 'Cause List']
                col_widths = [Cm(1.5), Cm(1.5), Cm(3.0), Cm(4.9), Cm(2.8), Cm(2.2), Cm(2.3)]
                for j, h in enumerate(headers):
                    set_cell_width(hdr[j], col_widths[j])
                    hdr[j].text = ''
                    p = hdr[j].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(h)
                    run.bold = True
                    set_run_font(run)
                    set_cell_shading(hdr[j], 'f5f5f5')

            _add_docx_data_row(table, entry, col_widths)
            i += 1

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="cause_list_{date_str}.docx"'
    doc.save(response)
    return response


@login_required
def cause_list_pdf(request):
    date_str = request.GET.get('date', '')
    if not date_str:
        return HttpResponse('No date provided.', status=400)

    try:
        date_obj = datetime.datetime.strptime(date_str, '%Y-%m-%d').date()
    except ValueError:
        return HttpResponse('Invalid date.', status=400)

    from weasyprint import HTML

    entries = CauseListEntry.objects.filter(date=date_obj).select_related('case')

    court_order = [
        'city_civil', 'small_causes', 'city_civil_rural', 'dist_session_rural',
        'prl_senior_rural', 'prl_junior_rural',
        'cmm', 'mmtc', 'mmtc_mayo', 'cjm_cmm_rural', 'vacation_court_rural',
        'family',
        'vacation_bench_family',
        'consumer', 'urban_consumer',
        'commercial', 'commercial_court_rural',
        'mayo_hall',
        'labour_court_rural', 'labour_court_urban', 'senior_anekal', 'junior_anekal', 'hosakote',
        'devanahalli', 'doddaballapur', 'nelamangala', 'kr_puram',
        'high_court_karnataka', 'supreme_court_india',
    ]
    diary_entries_for_stage = DiaryEntry.objects.filter(
        next_date=date_obj
    ).values('case_id', 'stage', 'entry_type', 'court', 'court_hall')
    stage_by_case = {}
    mediation_court_by_case = {}
    for de in diary_entries_for_stage:
        stage_by_case[de['case_id']] = de['stage']
        if de['entry_type'] == 'mediation':
            mc = de['court']
            mh = de['court_hall']
            if mc not in COURT_TO_BUILDING:
                mc = 'Karnataka Mediation Centre'
                mh = 'Mediation'
            mediation_court_by_case[de['case_id']] = (mc, mh)
    for e in entries:
        e.stage = stage_by_case.get(e.case.id, '')
        if e.case.id in mediation_court_by_case:
            e.mediation_court, e.mediation_court_hall = mediation_court_by_case[e.case.id]

    entries = sorted(entries, key=lambda e: (
        BUILDING_ORDER.index(COURT_TO_BUILDING.get(
            getattr(e, 'mediation_court', None) or e.case.court, 'other'))
            if (getattr(e, 'mediation_court', None) or e.case.court) in COURT_TO_BUILDING else 999,
        court_order.index(e.case.court) if e.case.court in court_order else 999,
        _tertiary_sort_key(e),
    ))
    for sl, e in enumerate(entries, 1):
        e.sl_no = sl

    from itertools import groupby
    def effective_building_code(e):
        court = getattr(e, 'mediation_court', None) or e.case.court
        return BUILDING_ORDER.index(COURT_TO_BUILDING.get(court, 'other')) if court in COURT_TO_BUILDING else 999
    building_groups = []
    for bldg_code, group in groupby(entries, key=effective_building_code):
        actual_code = BUILDING_ORDER[bldg_code] if bldg_code < len(BUILDING_ORDER) else 'other'
        building_groups.append((actual_code, list(group)))

    court_hall_incharges = set()
    for chi in CourtHallIncharge.objects.filter(date=date_obj, is_incharge=True):
        court_hall_incharges.add(f"{chi.court}__{chi.court_hall}")

    court_halls_on_date = []
    seen_halls = set()
    for e in entries:
        court = getattr(e, 'mediation_court', None) or e.case.court
        hall = getattr(e, 'mediation_court_hall', None) or e.case.court_hall
        key = (court, hall)
        if key not in seen_halls:
            seen_halls.add(key)
            court_halls_on_date.append(key)

    wednesdays, saturdays = _get_wed_sat_dates(date_obj)
    family_further_dates = _get_family_further_dates(date_obj)

    html_str = render(request, 'main/cause_list_pdf.html', {
        'entries': entries, 'building_groups': building_groups,
        'date_str': date_str, 'date_obj': date_obj,
        'court_labels': COURT_LABELS, 'court_to_building': COURT_TO_BUILDING,
        'court_hall_incharges': court_hall_incharges,
        'court_halls_on_date': court_halls_on_date,
        'wednesdays': wednesdays,
        'saturdays': saturdays,
        'family_further_dates': family_further_dates,
    }).content.decode()

    pdf_file = BytesIO()
    HTML(string=html_str).write_pdf(pdf_file)
    pdf_file.seek(0)

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="cause_list_{date_str}.pdf"'
    return response


# ── BATCH BUSINESS ENTRY ──

@login_required
def batch_new_case(request):
    cases = Case.objects.all().order_by('-id')
    selected_case = None
    case_id = request.GET.get('case_id') or request.POST.get('case_id')
    if case_id:
        selected_case = get_object_or_404(Case, id=case_id)

    if request.method == 'POST':
        case = get_object_or_404(Case, id=request.POST.get('case_id'))
        entries_data = []
        i = 0
        while True:
            previous_date = request.POST.get(f'previous_date_{i}')
            if previous_date is None:
                break
            next_date = request.POST.get(f'next_date_{i}')
            court = request.POST.get(f'court_{i}', '').strip()
            court_hall = request.POST.get(f'court_hall_{i}', '').strip()
            floor = request.POST.get(f'floor_{i}', '').strip()
            case_number_display = request.POST.get(f'case_number_display_{i}', '').strip()
            representing = request.POST.get(f'representing_{i}', '').strip()
            stage = request.POST.get(f'stage_{i}', '').strip()
            business = request.POST.get(f'business_{i}', '').strip()

            if previous_date and next_date and business:
                entries_data.append({
                    'previous_date': previous_date,
                    'next_date': next_date,
                    'court': court or COURT_LABELS.get(case.court, case.court),
                    'court_hall': court_hall or case.court_hall,
                    'floor': int(floor) if floor else case.floor,
                    'case_number_display': case_number_display or f"{case.case_type}/{case.case_number}/{case.case_year}",
                    'representing': representing or case.representing,
                    'stage': stage,
                    'business': business,
                })
            i += 1

        entries_data.sort(key=lambda e: e['previous_date'])

        for ed in entries_data:
            create_diary_entry(
                case=case,
                previous_date=ed['previous_date'],
                next_date=ed['next_date'],
                court=ed['court'],
                court_hall=ed['court_hall'],
                floor=ed['floor'],
                case_number_display=ed['case_number_display'],
                representing=ed['representing'],
                stage=ed['stage'],
                business=ed['business'],
                advocate=request.user,
            )

        if entries_data:
            messages.success(request, f'{len(entries_data)} business entr(y/ies) created.')
            return redirect('diary_entry_case', case_id=case.id)

    default_case_display = ''
    court_display = ''
    diary_entries = []
    if selected_case:
        default_case_display = f"{selected_case.case_type}/{selected_case.case_number}/{selected_case.case_year} — {selected_case.party_1} vs {selected_case.party_2}"
        court_display = COURT_LABELS.get(selected_case.court, selected_case.court)
        diary_entries = selected_case.diary_entries.all()

    return render(request, 'main/batch_new_case.html', {
        'cases': cases,
        'selected_case': selected_case,
        'default_case_display': default_case_display,
        'court_display': court_display,
        'court_labels': COURT_LABELS,
        'party1_choices': Party1Type.choices,
        'party2_choices': Party2Type.choices,
        'jurisdiction_choices': Jurisdiction.choices,
        'court_level_choices': CourtLevel.choices,
        'diary_entries': diary_entries,
    })


# ── TELEGRAM WEBHOOK ──

@csrf_exempt
@require_http_methods(['POST'])
def telegram_webhook(request):
    try:
        update = json.loads(request.body)
        update_id = update.get('update_id')
        if update_id:
            from .models import TelegramUpdate
            _, created = TelegramUpdate.objects.get_or_create(update_id=update_id)
            if not created:
                return HttpResponse('ok')
        from main.telegram_handler import process_update
        process_update(update)
    except Exception as e:
        logger.exception(f'Telegram webhook error: {e}')

    return HttpResponse('ok')


# ── COURT HALL SUGGESTIONS ──

@login_required
def suggest_court_halls(request):
    q = request.GET.get('q', '').strip()
    halls = Case.objects.values_list('court_hall', flat=True).distinct()
    if len(q) >= 1:
        halls = halls.filter(court_hall__icontains=q)
    halls = list(halls.order_by('court_hall')[:20])
    return JsonResponse([{'label': h, 'value': h} for h in halls], safe=False)


# ── COURT HALL NOTES ──

@login_required
def court_hall_notes(request):
    court = request.GET.get('court', '')
    court_hall = request.GET.get('court_hall', '')
    notes = CourtHallNote.objects.all()
    if court:
        notes = notes.filter(court=court)
    if court_hall:
        notes = notes.filter(court_hall__icontains=court_hall)
    return render(request, 'main/court_hall_notes.html', {
        'notes': notes.order_by('-updated_at'),
        'court': court,
        'court_hall': court_hall,
        'court_labels': COURT_LABELS,
        'court_halls': COURT_HALLS,
    })


@login_required
def add_court_hall_note(request):
    if request.method == 'POST':
        court_code = request.POST.get('court')
        court_hall = request.POST.get('court_hall')
        note = request.POST.get('note', '').strip()
        default_floor_raw = request.POST.get('default_floor', '').strip()
        if court_code and court_hall:
            obj, created = CourtHallNote.objects.get_or_create(
                court=court_code,
                court_hall=court_hall,
                defaults={'note': note},
            )
            if not created:
                if obj.note:
                    obj.note += f'\n\n---\n\n{note}'
                else:
                    obj.note = note
            if default_floor_raw:
                try:
                    obj.default_floor = int(default_floor_raw)
                except ValueError:
                    pass
            obj.save()
            messages.success(request, 'Court hall note saved.')
        else:
            messages.error(request, 'Court and Court Hall are required.')
        return redirect(request.POST.get('next', 'cause_list'))
    court = request.GET.get('court', '')
    court_hall = request.GET.get('court_hall', '')
    existing_note = None
    if court and court_hall:
        try:
            existing_note = CourtHallNote.objects.get(court=court, court_hall=court_hall)
        except CourtHallNote.DoesNotExist:
            pass
    return render(request, 'main/add_court_hall_note.html', {
        'court': court,
        'court_hall': court_hall,
        'next': request.GET.get('next', 'cause_list'),
        'court_labels': COURT_LABELS,
        'court_halls': COURT_HALLS,
        'existing_note': existing_note,
    })


# ── MEDIATION ──

@login_required
def add_mediation_business(request, case_id):
    case = get_object_or_404(Case, id=case_id)
    latest_data = get_latest_entry_data(case)

    if request.method == 'POST':
        previous_date = (request.POST.get('previous_date') or '').strip()
        stage = (request.POST.get('stage') or 'Mediation').strip()
        business = (request.POST.get('business') or '').strip()
        next_date = (request.POST.get('next_date') or '').strip()
        mediation_time = (request.POST.get('mediation_time') or '').strip()

        try:
            if mediation_time:
                parts = mediation_time.split(':')
                mediation_time_obj = datetime.time(int(parts[0]), int(parts[1]))
            else:
                mediation_time_obj = datetime.time(14, 30)
        except (ValueError, IndexError):
            mediation_time_obj = datetime.time(14, 30)

        entry = create_diary_entry(
            case=case, entry_type='mediation',
            previous_date=previous_date,
            court='Karnataka Mediation Centre',
            court_hall='Mediation',
            floor=0,
            case_number_display=f"{case.case_type}/{case.case_number}/{case.case_year}",
            representing=case.representing,
            representing_parties=case.representing_parties,
            party_1_total=case.party_1_total,
            party_2_total=case.party_2_total,
            stage=stage, business=business, next_date=next_date,
            mediation_time=mediation_time_obj,
            advocate=request.user,
        )

        messages.success(request, 'Mediation business added.')
        return redirect('diary_entry_case', case_id=case.id)

    return render(request, 'main/add_mediation_business.html', {
        'case': case, 'latest_data': latest_data,
        'today': datetime.date.today(),
        'default_time': '14:30',
    })


@login_required
def close_mediation(request, case_id):
    case = get_object_or_404(Case, id=case_id)

    if request.method == 'POST':
        status = request.POST.get('mediation_status')
        notes = request.POST.get('notes', '').strip()

        if status not in ('settled', 'failed'):
            messages.error(request, 'Invalid mediation status.')
            return redirect('diary_entry_case', case_id=case.id)

        case.mediation_status = status
        case.mediation_next_date = None
        case.save()

        label = dict(MediationStatus.choices).get(status, status)
        business_text = f'Mediation {label.lower()}.'
        if notes:
            business_text += f'\n\nNotes: {notes}'

        create_diary_entry(
            case=case, entry_type='mediation',
            previous_date=datetime.date.today(),
            court='Karnataka Mediation Centre',
            court_hall='Mediation', floor=0,
            case_number_display=f"{case.case_type}/{case.case_number}/{case.case_year}",
            representing=case.representing,
            representing_parties=case.representing_parties,
            party_1_total=case.party_1_total,
            party_2_total=case.party_2_total,
            stage='Mediation Closed',
            business=business_text,
            next_date=datetime.date.today(),
            advocate=request.user,
        )

        messages.success(request, f'Mediation closed — {label}.')
        return redirect('diary_entry_case', case_id=case.id)

    return render(request, 'main/close_mediation.html', {
        'case': case,
        'court_display': COURT_LABELS.get(case.court, case.court),
    })


# ── REMINDERS ──

@login_required
def mark_reminder_done(request, reminder_id):
    reminder = get_object_or_404(Reminder, id=reminder_id)
    reminder.completed = not reminder.completed
    reminder.save()
    return redirect('diary_entry_case', case_id=reminder.diary_entry.case.id)


@login_required
@user_passes_test(lambda u: u.is_superuser or u.userprofile.role == 'admin')
def send_reminders_now(request):
    from main.management.commands.send_reminders import send_due_reminders
    sent = send_due_reminders(auto=False)
    messages.success(request, f'{sent} reminder(s) sent to the Telegram group.')
    return redirect(request.META.get('HTTP_REFERER', 'home'))


# ── ECOURTS UPDATE ──

def _ecourts_update_access(user):
    if user.is_superuser:
        return True
    if hasattr(user, 'userprofile') and user.userprofile.role == UserRole.ADMIN:
        return True
    if hasattr(user, 'userprofile') and user.userprofile.can_access_ecourts:
        return True
    return SiteSetting.get_bool('ecourts_update_open', False)


@login_required
@user_passes_test(_ecourts_update_access)
def ecourts_update_list(request):
    """List cases and allow CNR entry + eCourts fetch trigger."""
    from django.db.models import Exists, OuterRef, Value, IntegerField
    from django.db.models import Case as DBCase, When

    # Only show cases from city_civil_complex, magistrates_complex, family_court
    FILTER_BUILDINGS = {'city_civil_complex', 'magistrates_complex', 'family_court'}
    MAGISTRATE_COURTS = {c for c, b in COURT_TO_BUILDING.items() if b == 'magistrates_complex'}
    ALLOWED_COURTS = [c for c, b in COURT_TO_BUILDING.items() if b in FILTER_BUILDINGS]

    q = request.GET.get('q', '').strip()

    has_ecourts_sub = DiaryEntry.objects.filter(
        case=OuterRef('pk'), entry_type='business'
    ).exclude(ecourts_business='').exclude(ecourts_business__isnull=True)

    cases = Case.objects.filter(
        court__in=ALLOWED_COURTS
    ).exclude(
        court__in=MAGISTRATE_COURTS, case_type__istartswith='cr'
    ).exclude(
        ecourts_status__in=['done', 'unsupported', 'no_data']
    ).annotate(
        has_ecourts=Exists(has_ecourts_sub),
        sort_order=DBCase(
            When(ecourts_status='pending', then=Value(2)),
            When(has_ecourts=True, then=Value(1)),
            default=Value(0),
            output_field=IntegerField(),
        )
    ).order_by('sort_order', '-id')

    if q:
        cases = cases.filter(
            Q(case_number__icontains=q) |
            Q(case_type__icontains=q) |
            Q(case_year__icontains=q) |
            Q(cnr__icontains=q) |
            Q(party_1__icontains=q) |
            Q(party_2__icontains=q)
        )

    paginator = Paginator(cases, 50)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)

    return render(request, 'main/ecourts_update.html', {
        'page_obj': page_obj,
        'cases': page_obj.object_list,
        'q': q,
        'is_admin': request.user.is_superuser or (hasattr(request.user, 'userprofile') and request.user.userprofile.role == UserRole.ADMIN),
        'ecourts_toggle_on': SiteSetting.get_bool('ecourts_update_open', False),
    })


@login_required
@user_passes_test(_ecourts_update_access)
def ecourts_update_single(request, case_id):
    """Handle POST: update CNR and/or trigger eCourts fetch for a case."""
    case = get_object_or_404(Case, id=case_id)

    if request.method == 'POST':
        if 'save_cnr' in request.POST:
            cnr = (request.POST.get('cnr', '') or '').strip()
            case.cnr = cnr
            if cnr:
                case.ecourts_status = 'pending'
                msg = f'CNR saved for {case.case_type}/{case.case_number}/{case.case_year}. Case queued for eCourts sync.'
            else:
                case.ecourts_status = ''
                msg = f'CNR cleared for {case.case_type}/{case.case_number}/{case.case_year}.'
            case.save()
            messages.success(request, msg)

        if 'fetch_ecourts' in request.POST:
            if not case.cnr:
                messages.error(request, 'No CNR set. Save a CNR first.')
            else:
                case.ecourts_status = 'pending'
                case.save()
                messages.success(request, f'Case queued for eCourts sync (status: pending). Run sync_ecourts.py on your laptop to fetch the data.')

    return redirect('ecourts_update_list')


@login_required
@require_http_methods(['POST'])
def ecourts_toggle_open(request):
    if not (request.user.is_superuser or (hasattr(request.user, 'userprofile') and request.user.userprofile.role == UserRole.ADMIN)):
        messages.error(request, 'Only admins can toggle this setting.')
        return redirect('ecourts_update_list')
    current = SiteSetting.get_bool('ecourts_update_open', False)
    SiteSetting.set_bool('ecourts_update_open', not current)
    messages.success(request, f'eCourts update page is now {"open" if not current else "closed"} to non-admin users.')
    return redirect('ecourts_update_list')


# ── HOME ──

@login_required
def home(request):
    return render(request, 'index.html')
